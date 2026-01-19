#!/usr/bin/env python3
"""
АСДУЕ - Веб-сканер сети с реальным сканированием
"""

from flask import Flask, render_template, request, jsonify, send_file
import json
import os
import threading
import time
from datetime import datetime
import csv
from docx import Document
import ipaddress

# Импортируем сканер правильно - из модуля scanner импортируем КЛАСС NetworkScanner
try:
    # Правильный импорт: из модуля scanner импортируем класс NetworkScanner
    from scanner import NetworkScanner  # Убедитесь, что в папке scanner есть файл __init__.py
    scanner = NetworkScanner()  # Создаем экземпляр класса
    print("✓ Сканер успешно импортирован")
except ImportError as e:
    print(f"✗ Ошибка импорта сканера: {e}")
    print("Проверьте структуру файлов:")
    print("  - scanner/__init__.py должен существовать")
    print("  - scanner/scanner.py должен содержать класс NetworkScanner")
    
    # Создаем заглушку для сканера
    class DummyScanner:
        def __init__(self):
            self.is_scanning = False
            
        def scan_network(self, network):
            print(f"Заглушка: сканирование сети {network}")
            return []
    
    scanner = DummyScanner()

app = Flask(__name__)

# Создаем директории
for folder in ['uploads', 'results', 'logs', 'static', 'templates']:
    os.makedirs(folder, exist_ok=True)

# Конфигурация
app.config['NETWORKS_FILE'] = 'networks.json'

# Глобальные переменные для статуса сканирования
scan_data = {
    'is_scanning': False,
    'progress': 0,
    'current_network': '',
    'total_networks': 0,
    'scanned_networks': 0,
    'hosts_found': 0,
    'results': [],
    'start_time': None,
    'end_time': None,
    'scan_thread': None,  # Этот объект не будет сериализован в JSON
    'logs': []
}

@app.route('/')
def index():
    """Главная страница АСДУЕ"""
    modules = [
        {'name': '📡 Сканирование сети', 'url': '/scan', 'desc': 'Запуск аудита сети', 'icon': 'radar'},
        {'name': '⚙️ Настройка сетей', 'url': '/networks', 'desc': 'Управление подсетями', 'icon': 'gear'},
        {'name': '📊 Результаты', 'url': '/results', 'desc': 'Просмотр и выгрузка', 'icon': 'table'},
        {'name': '📈 Статистика', 'url': '/stats', 'desc': 'Аналитика сети', 'icon': 'chart'},
    ]
    
    # Статистика
    stats = {
        'total_scans': len(scan_data['results']),
        'last_scan': scan_data['end_time'] if scan_data['end_time'] else 'Не выполнялось',
        'networks_count': len(load_networks())
    }
    
    return render_template('index.html', 
                         title="Аудит сети АСДУЕ",
                         modules=modules,
                         stats=stats,
                         is_scanning=scan_data['is_scanning'])

@app.route('/networks', methods=['GET', 'POST'])
def networks():
    """Управление сетями для сканирования"""
    
    if request.method == 'POST':
        action = request.form.get('action', '')
        
        if action == 'add':
            network = request.form.get('network', '').strip()
            if network:
                networks_list = load_networks()
                if network not in networks_list:
                    networks_list.append(network)
                    save_networks(networks_list)
        
        elif action == 'delete':
            network_to_delete = request.form.get('network_to_delete', '')
            if network_to_delete:
                networks_list = load_networks()
                if network_to_delete in networks_list:
                    networks_list.remove(network_to_delete)
                    save_networks(networks_list)
        
        elif action == 'clear':
            save_networks([])
        
        elif 'network_file' in request.files:
            file = request.files['network_file']
            if file.filename:
                content = file.read().decode('utf-8')
                new_networks = [line.strip() for line in content.split('\n') if line.strip()]
                networks_list = load_networks()
                networks_list.extend(new_networks)
                networks_list = list(set(networks_list))
                save_networks(networks_list)
    
    networks_list = load_networks()
    return render_template('networks.html', networks=networks_list)

@app.route('/scan', methods=['GET', 'POST'])
def scan_page():
    """Страница сканирования"""
    
    if request.method == 'POST':
        action = request.form.get('action', '')
        
        if action == 'start' and not scan_data['is_scanning']:
            # Запускаем сканирование в отдельном потоке
            scan_thread = threading.Thread(target=start_scanning, daemon=True)
            scan_thread.start()
            scan_data['scan_thread'] = scan_thread
            return jsonify({'status': 'started', 'message': 'Сканирование запущено'})
        
        elif action == 'stop' and scan_data['is_scanning']:
            scan_data['is_scanning'] = False
            return jsonify({'status': 'stopping', 'message': 'Остановка сканирования...'})
    
    networks_list = load_networks()
    return render_template('scan.html', 
                         networks=networks_list,
                         is_scanning=scan_data['is_scanning'],
                         scan_data=scan_data)

@app.route('/api/scan/status')
def scan_status():
    """API для получения статуса сканирования"""
    # Создаем копию данных без несериализуемых объектов
    serializable_data = {
        'is_scanning': scan_data['is_scanning'],
        'progress': scan_data['progress'],
        'current_network': scan_data['current_network'],
        'total_networks': scan_data['total_networks'],
        'scanned_networks': scan_data['scanned_networks'],
        'hosts_found': scan_data['hosts_found'],
        'results': scan_data['results'],
        'start_time': scan_data['start_time'],
        'end_time': scan_data['end_time'],
        'logs': scan_data['logs']
    }
    return jsonify(serializable_data)

@app.route('/api/scan/start', methods=['POST'])
def api_start_scan():
    """API для запуска сканирования"""
    if scan_data['is_scanning']:
        return jsonify({'status': 'error', 'message': 'Сканирование уже выполняется'})
    
    # Проверяем, есть ли сети для сканирования
    networks_list = load_networks()
    if not networks_list:
        return jsonify({'status': 'error', 'message': 'Нет сетей для сканирования'})
    
    scan_thread = threading.Thread(target=start_scanning, daemon=True)
    scan_thread.start()
    scan_data['scan_thread'] = scan_thread
    
    return jsonify({'status': 'success', 'message': 'Сканирование запущено'})

@app.route('/api/scan/stop', methods=['POST'])
def api_stop_scan():
    """API для остановки сканирования"""
    scan_data['is_scanning'] = False
    return jsonify({'status': 'success', 'message': 'Сканирование останавливается'})

@app.route('/api/networks/save', methods=['POST'])
def save_networks_api():
    """API для сохранения списка сетей"""
    try:
        data = request.get_json()
        networks = data.get('networks', [])
        
        # Базовая валидация
        valid_networks = []
        invalid_networks = []
        
        for network in networks:
            if validate_network(network):
                valid_networks.append(network)
            else:
                invalid_networks.append(network)
        
        if invalid_networks:
            print(f"Некорректные сети: {invalid_networks}")
        
        save_networks(valid_networks)
        
        return jsonify({
            'status': 'success',
            'message': f'Сохранено {len(valid_networks)} сетей',
            'saved_count': len(valid_networks),
            'invalid_networks': invalid_networks
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400

@app.route('/api/networks/list', methods=['GET'])
def list_networks_api():
    """API для получения списка сетей"""
    networks = load_networks()
    return jsonify({'networks': networks})

@app.route('/api/scan/logs', methods=['GET'])
def get_scan_logs():
    """API для получения логов сканирования"""
    return jsonify({'logs': scan_data.get('logs', [])})

@app.route('/results')
def results():
    """Страница с результатами сканирования"""
    return render_template('results.html', 
                         results=scan_data['results'],
                         hosts_count=len(scan_data['results']),
                         last_scan=scan_data['end_time'])

@app.route('/export/csv')
def export_csv():
    """Экспорт результатов в CSV"""
    if not scan_data['results']:
        return jsonify({'status': 'error', 'message': 'Нет данных для экспорта'}), 400
    
    filename = f"scan_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    filepath = os.path.join('results', filename)
    
    # Определяем все возможные ключи
    all_keys = set()
    for host in scan_data['results']:
        all_keys.update(host.keys())
    
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=sorted(all_keys))
        writer.writeheader()
        writer.writerows(scan_data['results'])
    
    return send_file(filepath, as_attachment=True)

@app.route('/export/docx')
def export_docx():
    """Экспорт результатов в Word"""
    if not scan_data['results']:
        return jsonify({'status': 'error', 'message': 'Нет данных для экспорта'}), 400
    
    filename = f"scan_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join('results', filename)
    
    doc = Document()
    
    # Заголовок
    doc.add_heading('Результаты сканирования сети АСДУЕ', 0)
    
    # Общая информация
    doc.add_paragraph(f'Дата сканирования: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Найдено устройств: {len(scan_data["results"])}')
    doc.add_paragraph(f'Время начала: {scan_data.get("start_time", "N/A")}')
    doc.add_paragraph(f'Время окончания: {scan_data.get("end_time", "N/A")}')
    
    doc.add_page_break()
    
    # Таблица с результатами
    doc.add_heading('Обнаруженные устройства', level=1)
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки таблицы
    headers = table.rows[0].cells
    headers[0].text = 'IP адрес'
    headers[1].text = 'Имя хоста'
    headers[2].text = 'MAC адрес'
    headers[3].text = 'Производитель'
    headers[4].text = 'ОС'
    headers[5].text = 'Статус'
    headers[6].text = 'Время сканирования'
    
    # Данные
    for host in scan_data['results']:
        row_cells = table.add_row().cells
        row_cells[0].text = host.get('ip', '')
        row_cells[1].text = host.get('hostname', '')
        row_cells[2].text = host.get('mac', '')
        row_cells[3].text = host.get('vendor', '')
        row_cells[4].text = host.get('os', '')
        row_cells[5].text = host.get('status', '')
        row_cells[6].text = host.get('scan_time', '')
    
    doc.save(filepath)
    return send_file(filepath, as_attachment=True)

@app.route('/stats')
def stats():
    """Страница статистики"""
    stats_data = {
        'total_hosts': len(scan_data['results']),
        'vendors': {},
        'os_distribution': {},
        'scan_history': []
    }
    
    # Анализ производителей
    for host in scan_data['results']:
        vendor = host.get('vendor', 'Unknown')
        stats_data['vendors'][vendor] = stats_data['vendors'].get(vendor, 0) + 1
        
        os_name = host.get('os', 'Unknown')
        stats_data['os_distribution'][os_name] = stats_data['os_distribution'].get(os_name, 0) + 1
    
    return render_template('stats.html', stats=stats_data)

@app.route('/health')
def health():
    """Проверка здоровья сервиса"""
    return jsonify({
        'status': 'healthy',
        'service': 'network-audit-asdue',
        'version': '1.0.0',
        'timestamp': datetime.now().isoformat(),
        'scanning': scan_data['is_scanning'],
        'hosts_in_memory': len(scan_data['results'])
    })

def load_networks():
    """Загрузка списка сетей"""
    networks_file = app.config['NETWORKS_FILE']
    if os.path.exists(networks_file):
        try:
            with open(networks_file, 'r') as f:
                return json.load(f)
        except:
            return []
    return []

def save_networks(networks):
    """Сохранение списка сетей"""
    with open(app.config['NETWORKS_FILE'], 'w') as f:
        json.dump(networks, f, indent=2)

def validate_network(network_str):
    """Проверка корректности формата сети CIDR"""
    try:
        ipaddress.ip_network(network_str, strict=False)
        return True
    except:
        return False

def add_scan_log(message, log_type='info'):
    """Добавление записи в логи сканирования"""
    timestamp = datetime.now().strftime('%H:%M:%S')
    log_entry = {
        'timestamp': timestamp,
        'message': message,
        'type': log_type
    }
    scan_data['logs'].append(log_entry)
    
    # Ограничиваем количество логов (последние 100)
    if len(scan_data['logs']) > 100:
        scan_data['logs'] = scan_data['logs'][-100:]

def start_scanning():
    """Функция запуска сканирования (работает в отдельном потоке)"""
    global scan_data
    
    # Сброс предыдущих данных
    scan_data['is_scanning'] = True
    scan_data['progress'] = 0
    scan_data['results'] = []
    scan_data['current_network'] = ''
    scan_data['start_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    scan_data['end_time'] = None
    scan_data['hosts_found'] = 0
    scan_data['logs'] = []
    
    networks_list = load_networks()
    scan_data['total_networks'] = len(networks_list)
    scan_data['scanned_networks'] = 0
    
    add_scan_log('Начинаем сканирование...', 'info')
    
    if not networks_list:
        add_scan_log('Нет сетей для сканирования!', 'error')
        print("Нет сетей для сканирования!")
        scan_data['is_scanning'] = False
        return
    
    add_scan_log(f'Начинаем сканирование {len(networks_list)} сетей...', 'info')
    print(f"Начинаем сканирование {len(networks_list)} сетей...")
    
    # Сканируем каждую сеть
    all_results = []
    
    for i, network in enumerate(networks_list):
        if not scan_data['is_scanning']:
            add_scan_log('Сканирование прервано пользователем', 'warning')
            print("Сканирование прервано пользователем")
            break
        
        scan_data['current_network'] = network
        scan_data['scanned_networks'] = i + 1
        scan_data['progress'] = int((i + 1) / len(networks_list) * 100)
        
        log_msg = f"Сканируем сеть {i+1}/{len(networks_list)}: {network}"
        add_scan_log(log_msg, 'info')
        print(log_msg)
        
        try:
            # Используем наш сканер
            network_results = scanner.scan_network(network)
            
            all_results.extend(network_results)
            scan_data['results'] = all_results
            scan_data['hosts_found'] = len(all_results)
            
            add_scan_log(f"Найдено устройств в сети {network}: {len(network_results)}", 'success')
            add_scan_log(f"Всего найдено: {len(all_results)} устройств", 'info')
            
            print(f"  Найдено устройств в этой сети: {len(network_results)}")
            print(f"  Всего найдено: {len(all_results)} устройств")
            
        except Exception as e:
            error_msg = f"Ошибка при сканировании сети {network}: {e}"
            add_scan_log(error_msg, 'error')
            print(error_msg)
        
        # Небольшая пауза между сетями
        time.sleep(1)
    
    # Завершение сканирования
    scan_data['is_scanning'] = False
    scan_data['end_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    scan_data['progress'] = 100
    
    completion_msg = f"Сканирование завершено! Найдено устройств: {len(all_results)}"
    add_scan_log(completion_msg, 'success')
    print(completion_msg)
    
    # Сохраняем результаты в файл
    if all_results:
        save_results_to_file(all_results)

def save_results_to_file(results):
    """Сохранение результатов в JSON файл"""
    filename = f"scan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    filepath = os.path.join('results', filename)
    
    try:
        with open(filepath, 'w') as f:
            json.dump({
                'scan_time': datetime.now().isoformat(),
                'total_hosts': len(results),
                'hosts': results
            }, f, indent=2, default=str)
        
        add_scan_log(f"Результаты сохранены в {filepath}", 'info')
        print(f"Результаты сохранены в {filepath}")
    except Exception as e:
        error_msg = f"Ошибка сохранения результатов: {e}"
        add_scan_log(error_msg, 'error')
        print(f"Ошибка сохранения результатов: {e}")

if __name__ == '__main__':
    # Создаем networks.json если его нет
    if not os.path.exists('networks.json'):
        with open('networks.json', 'w') as f:
            json.dump(["192.168.1.0/24", "10.0.0.0/24"], f, indent=2)
    
    print("=" * 60)
    print("АСДУЕ - Веб-сканер сети")
    print("=" * 60)
    print(f"Сервер запускается на http://127.0.0.1:5000")
    print("")
    print("Проверьте настройки:")
    print(f"  - networks.json содержит {len(load_networks())} сетей")
    
    # Проверяем доступность nmap
    try:
        import nmap
        print("  - nmap доступен: Да")
    except ImportError:
        print("  - nmap доступен: Нет (установите: pip install python-nmap)")
    
    # Проверяем, что сканер работает
    try:
        test_result = scanner.scan_network("127.0.0.1/32")
        print(f"  - сканер работает: Да (тест: {len(test_result)} устройств)")
    except Exception as e:
        print(f"  - сканер работает: Нет (ошибка: {e})")
    
    print("=" * 60)
    
    app.run(debug=True, host='0.0.0.0')